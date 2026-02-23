"""
Generador de documento Word con observaciones y recomendaciones de distribución.
"""
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def _style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'


def _header_row(table, headers):
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1F4E79',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)


def generate_word(result, linea, fecha, tallas):
    """Genera documento Word con observaciones de distribución.

    Args:
        result: dict retornado por run_distribution()
        linea: "DAMA" o "CABALLERO"
        fecha: string con fecha
        tallas: lista de tallas

    Returns:
        BytesIO con el documento Word
    """
    distribuciones = result['distribuciones']
    summary = result['summary']
    resumen_tienda = summary['resumen_tienda']
    resumen_prioridad = summary['resumen_prioridad']
    resumen_sublinea = summary['resumen_sublinea']

    doc = Document()

    # Ajustar márgenes
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Título
    title = doc.add_heading(f'Distribución desde Almacén - {linea}', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph(f'Fecha: {fecha}')

    # --- RESUMEN EJECUTIVO ---
    _add_heading(doc, 'Resumen Ejecutivo')

    total_pares = summary['total_pares']
    total_almacen = summary['total_almacen']
    pct = summary['pct_distribuido']
    total_valor = summary['total_valor']
    n_tiendas = len(resumen_tienda)
    n_productos = summary['total_productos']
    restante = summary['total_restante']

    p = doc.add_paragraph()
    p.add_run('Stock en Almacén: ').bold = True
    p.add_run(f'{total_almacen:,} pares en {summary["productos_almacen"]} productos')
    p = doc.add_paragraph()
    p.add_run('Distribución propuesta: ').bold = True
    p.add_run(f'{total_pares:,} pares ({pct:.0f}% del almacén) → {n_tiendas} tiendas')
    p = doc.add_paragraph()
    p.add_run('Productos incluidos: ').bold = True
    p.add_run(f'{n_productos}')
    p = doc.add_paragraph()
    p.add_run('Valor estimado: ').bold = True
    p.add_run(f'${total_valor:,.0f} MXN')
    p = doc.add_paragraph()
    p.add_run('Restante en almacén: ').bold = True
    p.add_run(f'{restante:,} pares')

    # --- DISTRIBUCIÓN POR PRIORIDAD ---
    _add_heading(doc, 'Distribución por Prioridad')

    prioridades = ['URGENTE', 'ALTA', 'NORMAL', 'BAJA', 'OPORTUNIDAD']
    table = doc.add_table(rows=1, cols=3)
    _header_row(table, ['Prioridad', 'Pares', '% del Total'])

    for prio in prioridades:
        qty = resumen_prioridad.get(prio, 0)
        if qty > 0:
            row = table.add_row()
            row.cells[0].text = prio
            row.cells[1].text = f'{qty:,}'
            row.cells[2].text = f'{qty / total_pares * 100:.1f}%' if total_pares > 0 else '0%'

    row = table.add_row()
    row.cells[0].text = 'TOTAL'
    row.cells[1].text = f'{total_pares:,}'
    row.cells[2].text = '100%'
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    _style_table(table)

    # --- DISTRIBUCIÓN POR CATEGORÍA ---
    _add_heading(doc, 'Distribución por Categoría')

    table = doc.add_table(rows=1, cols=3)
    _header_row(table, ['Sublínea', 'Pares', '% del Total'])

    for sub, qty in sorted(resumen_sublinea.items(), key=lambda x: -x[1]):
        row = table.add_row()
        row.cells[0].text = sub
        row.cells[1].text = f'{qty:,}'
        row.cells[2].text = f'{qty / total_pares * 100:.1f}%' if total_pares > 0 else '0%'

    _style_table(table)

    # --- RESUMEN POR TIENDA ---
    _add_heading(doc, 'Resumen por Tienda')

    table = doc.add_table(rows=1, cols=6)
    _header_row(table, ['Tienda', 'Pares', 'Productos', 'Urgente', 'Alta', 'Valor $'])

    for t in sorted(resumen_tienda.keys()):
        r = resumen_tienda[t]
        row = table.add_row()
        row.cells[0].text = f'Tienda {t}'
        row.cells[1].text = f'{r["pares"]:,}'
        row.cells[2].text = str(len(r['productos']))
        row.cells[3].text = str(r['urgente'])
        row.cells[4].text = str(r['alta'])
        row.cells[5].text = f'${r["valor"]:,.0f}'

    _style_table(table)

    # --- RECOMENDACIONES ---
    _add_heading(doc, 'Observaciones y Recomendaciones')

    # Tiendas con más urgencia
    tiendas_urgentes = [(t, r) for t, r in resumen_tienda.items() if r['urgente'] > 0]
    tiendas_urgentes.sort(key=lambda x: -x[1]['urgente'])

    if tiendas_urgentes:
        doc.add_paragraph(
            'Tiendas con productos URGENTES (menos de 5 días de inventario):',
            style='List Bullet'
        )
        for t, r in tiendas_urgentes[:5]:
            doc.add_paragraph(
                f'Tienda {t}: {r["urgente"]} pares urgentes de {len(r["productos"])} productos',
                style='List Bullet 2'
            )

    # Tiendas con más volumen
    tiendas_vol = sorted(resumen_tienda.items(), key=lambda x: -x[1]['pares'])
    doc.add_paragraph(
        'Tiendas con mayor volumen de distribución:',
        style='List Bullet'
    )
    for t, r in tiendas_vol[:5]:
        doc.add_paragraph(
            f'Tienda {t}: {r["pares"]:,} pares ({len(r["productos"])} productos)',
            style='List Bullet 2'
        )

    # Nota general
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Nota: ').bold = True
    p.add_run(
        'Esta distribución fue generada por el sistema automatizado basado en '
        'el algoritmo de scoring que considera: ventas de los últimos 15 días (60%), '
        'peso general de cada tienda (30%), y factor de urgencia por nivel de inventario. '
        'Se recomienda revisar las picking lists por tienda antes de ejecutar el envío.'
    )

    # Guardar a BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
